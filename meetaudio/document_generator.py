"""
WORD文档生成器
"""

import os
import io
from typing import Dict, Any, Optional
from datetime import datetime
import logging

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """WORD文档生成器"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            logger.warning("python-docx未安装，WORD文档生成功能不可用")
    
    def generate_meeting_minutes_doc(
        self,
        minutes_data: Dict[str, Any],
        template_style: str = "standard"
    ) -> Optional[bytes]:
        """
        生成会议纪要WORD文档
        
        Args:
            minutes_data: 会议纪要数据
            template_style: 模板样式
            
        Returns:
            WORD文档字节流
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx未安装，无法生成WORD文档")
            return None
        
        try:
            doc = Document()
            
            # 设置文档样式
            self._setup_document_styles(doc)
            
            # 添加标题
            self._add_title(doc, minutes_data.get("title", "会议纪要"))
            
            # 添加会议基本信息
            self._add_meeting_header(doc, minutes_data.get("header", {}))
            
            # 添加会议内容
            self._add_meeting_content(doc, minutes_data.get("content", {}))
            
            # 添加页脚
            self._add_meeting_footer(doc, minutes_data.get("footer", {}))
            
            # 保存到字节流
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            logger.info("WORD文档生成成功")
            return doc_bytes.getvalue()
            
        except Exception as e:
            logger.error(f"生成WORD文档失败: {e}")
            return None
    
    def _setup_document_styles(self, doc):
        """设置文档样式"""
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5  # 1.5倍行距
        style.paragraph_format.space_after = Pt(6)

        # 创建标题样式
        if 'Meeting Title' not in [s.name for s in doc.styles]:
            title_style = doc.styles.add_style('Meeting Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = '黑体'
            title_font.size = Pt(18)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_before = Pt(0)
            title_style.paragraph_format.space_after = Pt(18)
            title_style.paragraph_format.line_spacing = 1.0

        # 创建一级标题样式
        if 'Heading 1 Custom' not in [s.name for s in doc.styles]:
            h1_style = doc.styles.add_style('Heading 1 Custom', WD_STYLE_TYPE.PARAGRAPH)
            h1_font = h1_style.font
            h1_font.name = '黑体'
            h1_font.size = Pt(14)
            h1_font.bold = True
            h1_style.paragraph_format.space_before = Pt(12)
            h1_style.paragraph_format.space_after = Pt(6)
            h1_style.paragraph_format.line_spacing = 1.0

        # 创建正文样式
        if 'Body Text' not in [s.name for s in doc.styles]:
            body_style = doc.styles.add_style('Body Text', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = '宋体'
            body_font.size = Pt(12)
            body_style.paragraph_format.first_line_indent = Inches(0.25)
            body_style.paragraph_format.line_spacing = 1.5
            body_style.paragraph_format.space_after = Pt(6)

        # 创建信息项样式
        if 'Info Item' not in [s.name for s in doc.styles]:
            info_style = doc.styles.add_style('Info Item', WD_STYLE_TYPE.PARAGRAPH)
            info_font = info_style.font
            info_font.name = '宋体'
            info_font.size = Pt(12)
            info_style.paragraph_format.left_indent = Inches(0.5)
            info_style.paragraph_format.line_spacing = 1.5
            info_style.paragraph_format.space_after = Pt(3)
    
    def _add_title(self, doc, title: str):
        """添加标题"""
        title_para = doc.add_paragraph(title)
        title_para.style = 'Meeting Title'
    
    def _add_meeting_header(self, doc, header_data: Dict[str, Any]):
        """添加会议基本信息"""
        doc.add_paragraph()  # 空行

        # 会议基本信息
        info_items = [
            ("会议名称", header_data.get("meeting_name", "")),
            ("会议时间", header_data.get("date", "")),
            ("会议地点", header_data.get("location", "")),
            ("主持人", header_data.get("host", "")),
            ("记录人", header_data.get("recorder", "")),
        ]

        for label, value in info_items:
            if value:
                para = doc.add_paragraph(f"{label}：{value}")
                para.style = 'Info Item'

        # 参会人员
        attendees = header_data.get("attendees", [])
        if attendees:
            if isinstance(attendees, list):
                attendees_text = "、".join(attendees)
            else:
                attendees_text = str(attendees)

            para = doc.add_paragraph(f"参会人员：{attendees_text}")
            para.style = 'Info Item'

        doc.add_paragraph()  # 空行
    
    def _add_meeting_content(self, doc, content_data: Dict[str, Any]):
        """添加会议内容"""
        
        # 会议概况
        summary = content_data.get("summary", "")
        if summary:
            self._add_section(doc, "一、会议概况", summary)
        
        # 主要决策事项
        decisions = content_data.get("decisions", [])
        if decisions:
            decisions_text = self._format_list_items(decisions)
            self._add_section(doc, "二、主要决策事项", decisions_text)
        
        # 工作安排
        action_items = content_data.get("action_items", [])
        if action_items:
            actions_text = self._format_list_items(action_items)
            self._add_section(doc, "三、工作安排", actions_text)
        
        # 领导讲话要点
        leadership_remarks = content_data.get("leadership_remarks", {})
        if leadership_remarks:
            remarks_text = self._format_leadership_remarks(leadership_remarks)
            self._add_section(doc, "四、领导讲话要点", remarks_text)
        
        # 下一步工作
        next_steps = content_data.get("next_steps", [])
        if next_steps:
            next_steps_text = self._format_list_items(next_steps)
            self._add_section(doc, "五、下一步工作", next_steps_text)
    
    def _add_section(self, doc, title: str, content: str):
        """添加章节"""
        # 添加标题
        title_para = doc.add_paragraph(title)
        title_para.style = 'Heading 1 Custom'

        # 添加内容
        if content.strip():
            # 处理多行内容
            lines = content.strip().split('\n')
            for line in lines:
                if line.strip():
                    content_para = doc.add_paragraph(line.strip())
                    content_para.style = 'Body Text'

        doc.add_paragraph()  # 空行
    
    def _format_list_items(self, items: list) -> str:
        """格式化列表项"""
        if not items:
            return ""

        formatted_items = []
        for i, item in enumerate(items, 1):
            if isinstance(item, str) and item.strip():
                # 清理多余的空格和换行
                clean_item = ' '.join(item.strip().split())
                formatted_items.append(f"{i}. {clean_item}")

        return "\n".join(formatted_items)
    
    def _format_leadership_remarks(self, remarks: Dict[str, str]) -> str:
        """格式化领导讲话"""
        if not remarks:
            return ""

        formatted_remarks = []
        for leader, remark in remarks.items():
            if remark.strip():
                # 清理多余的空格和换行
                clean_remark = ' '.join(remark.strip().split())
                formatted_remarks.append(f"{leader}：{clean_remark}")

        return "\n\n".join(formatted_remarks)
    
    def _add_meeting_footer(self, doc, footer_data: Dict[str, Any]):
        """添加页脚信息"""
        doc.add_paragraph()  # 空行
        doc.add_paragraph()  # 额外空行

        # 确认信息
        confirm_para = doc.add_paragraph("本纪要已经与会人员确认。")
        confirm_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        confirm_para.paragraph_format.space_after = Pt(12)

        # 记录人和日期
        recorder = footer_data.get("recorder", "")
        review_date = footer_data.get("review_date", datetime.now().strftime("%Y年%m月%d日"))

        if recorder:
            recorder_para = doc.add_paragraph(f"记录人：{recorder}")
            recorder_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            recorder_para.paragraph_format.space_after = Pt(6)

        date_para = doc.add_paragraph(f"日期：{review_date}")
        date_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def generate_simple_doc(self, title: str, content: str) -> Optional[bytes]:
        """
        生成简单文档
        
        Args:
            title: 文档标题
            content: 文档内容
            
        Returns:
            WORD文档字节流
        """
        if not DOCX_AVAILABLE:
            return None
        
        try:
            doc = Document()
            
            # 添加标题
            title_para = doc.add_paragraph(title)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.runs[0]
            title_run.font.size = Pt(16)
            title_run.font.bold = True
            
            doc.add_paragraph()  # 空行
            
            # 添加内容
            content_para = doc.add_paragraph(content)
            content_para.paragraph_format.first_line_indent = Inches(0.25)
            
            # 保存到字节流
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            return doc_bytes.getvalue()
            
        except Exception as e:
            logger.error(f"生成简单文档失败: {e}")
            return None


# 全局实例
document_generator = DocumentGenerator()
