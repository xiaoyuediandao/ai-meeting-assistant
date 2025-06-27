"""
AI撰稿引擎 - 基于豆包大模型的会议纪要生成
"""

import os
import json
import logging
import time
import threading
import io
from typing import Dict, List, Any, Optional
from datetime import datetime
import httpx
from openai import OpenAI
from .enhanced_client import MeetingResult
from .aviation_terms import aviation_processor

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available, Word export will be disabled")

logger = logging.getLogger(__name__)


class AIWriter:
    """AI撰稿引擎"""

    def __init__(self, api_key=None, model=None, base_url=None, timeout=None):
        self.templates = MeetingTemplates()
        self.client = None
        self.api_key = api_key or os.getenv("ARK_API_KEY")
        self.model = model or os.getenv("ARK_MODEL", "ARK_MODEL_EP")
        self.base_url = base_url or os.getenv("ARK_BASE_URL", "https://ark.cn-beijing.volces.com/api/v3")
        self.timeout = timeout or int(os.getenv("ARK_TIMEOUT", "300"))
        self._init_ai_client()

    def _init_ai_client(self):
        """初始化AI客户端"""
        try:
            if self.api_key:
                # 创建自定义的httpx客户端，设置充足的超时时间
                custom_timeout = httpx.Timeout(
                    timeout=float(self.timeout),   # 使用配置的超时时间
                    connect=30.0,    # 连接超时30秒
                    read=float(self.timeout) - 30.0,      # 读取超时
                    write=30.0       # 写入超时30秒
                )

                self.client = OpenAI(
                    base_url=self.base_url,
                    api_key=self.api_key,
                    timeout=custom_timeout,
                    max_retries=1  # 允许1次重试
                )
                logger.info("豆包AI客户端初始化成功")
            else:
                logger.warning("ARK_API_KEY未设置，将使用模拟AI生成")
        except Exception as e:
            logger.error(f"AI客户端初始化失败: {e}")
            self.client = None
        
    def generate_meeting_minutes(
        self,
        meeting_result: MeetingResult,
        meeting_info: Dict[str, Any],
        focus_on_last_speakers: bool = True,
        speaker_count: int = 2
    ) -> Dict[str, Any]:
        """
        生成会议纪要
        
        Args:
            meeting_result: 会议识别结果
            meeting_info: 会议基本信息
            focus_on_last_speakers: 是否聚焦最后几位发言人
            speaker_count: 聚焦的发言人数量
            
        Returns:
            会议纪要内容
        """
        logger.info("开始生成会议纪要")
        
        # 1. 预处理文本
        processed_text = self._preprocess_text(meeting_result.full_text)
        
        # 2. 提取关键信息
        key_info = meeting_result.extract_key_information()
        
        # 3. 获取重点发言内容
        if focus_on_last_speakers:
            last_speakers = meeting_result.get_last_speakers(speaker_count)
            focus_content = self._format_speaker_content(last_speakers)
        else:
            focus_content = processed_text
        
        # 4. 生成纪要内容
        minutes_content = self._generate_content(
            focus_content, 
            key_info, 
            meeting_info
        )
        
        # 5. 格式化输出
        formatted_minutes = self._format_minutes(minutes_content, meeting_info)
        
        logger.info("会议纪要生成完成")
        return formatted_minutes
    
    def _preprocess_text(self, text: str) -> str:
        """预处理文本"""
        # 使用民航术语处理器规范化
        normalized_text = aviation_processor.normalize_text(text)
        
        # 清理无意义的口语化表达
        cleaned_text = self._clean_filler_words(normalized_text)
        
        return cleaned_text
    
    def _clean_filler_words(self, text: str) -> str:
        """清理口语化填充词"""
        filler_words = [
            "嗯", "啊", "呃", "那个", "这个", "就是说", "然后呢", 
            "对吧", "是不是", "怎么说呢", "应该说", "可以说"
        ]
        
        result = text
        for word in filler_words:
            result = result.replace(word, "")
        
        # 清理多余空格
        import re
        result = re.sub(r'\s+', ' ', result).strip()
        
        return result
    
    def _format_speaker_content(self, speakers: List[tuple]) -> str:
        """格式化说话人内容"""
        if not speakers:
            return ""
        
        formatted_content = []
        
        # 假设最后两位是党委书记和总经理
        role_mapping = {
            0: "党委书记",
            1: "总经理"
        }
        
        for i, (speaker_id, content) in enumerate(speakers):
            role = role_mapping.get(i, f"发言人{i+1}")
            formatted_content.append(f"\n{role}发言：\n{content}")
        
        return "\n".join(formatted_content)
    
    def _generate_content(
        self,
        content: str,
        key_info: Dict[str, Any],
        meeting_info: Dict[str, Any]
    ) -> Dict[str, Any]:
        """生成纪要内容"""

        if self.client:
            # 使用豆包AI生成
            return self._ai_generate_content(content, key_info, meeting_info)
        else:
            # 使用规则生成示例
            generated_content = {
                "meeting_summary": self._generate_summary(content, meeting_info),
                "key_decisions": key_info.get("decisions", []),
                "action_items": key_info.get("actions", []),
                "next_steps": self._extract_next_steps(content),
                "leadership_remarks": self._extract_leadership_remarks(content),
            }
            return generated_content

    def _ai_generate_content(
        self,
        content: str,
        key_info: Dict[str, Any],
        meeting_info: Dict[str, Any]
    ) -> Dict[str, Any]:
        """使用豆包AI生成会议纪要内容"""

        # 检查是否应该跳过AI生成（仅用于调试）
        skip_ai = os.getenv("SKIP_AI_GENERATION", "false").lower() == "true"

        if skip_ai:
            logger.info("调试模式：跳过AI生成，使用本地生成")
            return self._generate_local_content(content, key_info, meeting_info)

        try:
            # 使用之前能工作的模型
            model = os.getenv("ARK_MODEL", "ep-20250511141457-zf4l4")

            # 检查内容长度，决定是否需要分段处理
            max_content_length = 8000  # 单次处理的最大字符数
            if len(content) > max_content_length:
                logger.info(f"内容较长（{len(content)}字符），使用分段处理")
                return self._chunked_ai_generate(content, key_info, meeting_info, model)

            # 构建系统提示词
            system_prompt = self._build_system_prompt()

            # 构建用户提示词
            user_prompt = self._build_user_prompt(content, key_info, meeting_info)

            logger.info(f"调用豆包AI生成会议纪要，内容长度: {len(content)}字符")

            # 记录开始时间
            start_time = time.time()

            # 尝试直接使用httpx调用，避免OpenAI客户端的问题
            try:
                ai_response = self._direct_api_call(model, system_prompt, user_prompt)
            except Exception as direct_error:
                logger.warning(f"直接API调用失败: {direct_error}，尝试OpenAI客户端")
                # 降级到OpenAI客户端
                response = self.client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=0.3,  # 降低随机性，提高一致性
                    max_tokens=3000,  # 恢复合理的生成长度
                    timeout=1800  # 30分钟超时
                )
                ai_response = response.choices[0].message.content

            # 记录结束时间
            end_time = time.time()
            duration = end_time - start_time
            logger.info(f"豆包AI调用耗时: {duration:.2f}秒")

            logger.info("豆包AI生成会议纪要成功")
            return self._parse_ai_response(ai_response)

        except Exception as e:
            error_msg = str(e)
            logger.error(f"AI生成失败，错误详情: {error_msg}")

            # 根据错误类型提供更具体的日志
            if "timeout" in error_msg.lower():
                logger.error("AI API调用超时，可能是网络问题或服务响应慢")
            elif "connection" in error_msg.lower():
                logger.error("AI API连接失败，请检查网络连接和API配置")
            elif "401" in error_msg or "unauthorized" in error_msg.lower():
                logger.error("AI API认证失败，请检查ARK_API_KEY配置")
            elif "429" in error_msg or "rate limit" in error_msg.lower():
                logger.error("AI API调用频率超限，请稍后重试")

            # 降级到本地生成
            logger.info("使用本地规则生成会议纪要")
            return self._generate_local_content(content, key_info, meeting_info)

    def _direct_api_call(self, model: str, system_prompt: str, user_prompt: str) -> str:
        """直接使用httpx调用豆包API，带重试机制"""
        import time

        ark_api_key = os.getenv("ARK_API_KEY")
        ark_base_url = os.getenv("ARK_BASE_URL", "https://ark.cn-beijing.volces.com/api/v3")

        url = f"{ark_base_url}/chat/completions"
        headers = {
            "Authorization": f"Bearer {ark_api_key}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "temperature": 0.3,
            "max_tokens": 3000
        }

        # 重试配置
        max_retries = 3
        base_delay = 5  # 基础延迟5秒

        for attempt in range(max_retries):
            try:
                logger.info(f"豆包AI调用尝试 {attempt + 1}/{max_retries}")

                # 使用合理的超时时间
                timeout = httpx.Timeout(
                    timeout=600.0,    # 总超时600秒（10分钟）
                    connect=30.0,     # 连接超时30秒
                    read=570.0,       # 读取超时570秒
                    write=30.0        # 写入超时30秒
                )

                with httpx.Client(timeout=timeout) as client:
                    response = client.post(url, headers=headers, json=payload)
                    response.raise_for_status()

                    result = response.json()
                    if "choices" in result and len(result["choices"]) > 0:
                        content = result["choices"][0]["message"]["content"]
                        if content and content.strip():
                            logger.info(f"豆包AI调用成功，尝试次数: {attempt + 1}")
                            return content
                        else:
                            raise Exception("API返回空内容")
                    else:
                        raise Exception(f"API响应格式错误: {result}")

            except Exception as e:
                error_msg = str(e)
                logger.warning(f"豆包AI调用失败 (尝试 {attempt + 1}/{max_retries}): {error_msg}")

                # 如果是最后一次尝试，抛出异常
                if attempt == max_retries - 1:
                    logger.error(f"豆包AI调用最终失败，已重试 {max_retries} 次")
                    raise Exception(f"豆包AI调用失败，已重试{max_retries}次: {error_msg}")

                # 指数退避延迟
                delay = base_delay * (2 ** attempt)
                logger.info(f"等待 {delay} 秒后重试...")
                time.sleep(delay)

    def _generate_local_content(self, content: str, key_info: Dict[str, Any], meeting_info: Dict[str, Any]) -> Dict[str, Any]:
        """本地生成会议纪要内容"""
        return {
            "meeting_summary": self._generate_summary(content, meeting_info),
            "key_decisions": key_info.get("decisions", []),
            "action_items": key_info.get("actions", []),
            "next_steps": self._extract_next_steps(content),
            "leadership_remarks": self._extract_leadership_remarks(content),
        }

    def _build_system_prompt(self) -> str:
        """构建系统提示词"""
        return """
你的任务是从会议录音中提取核心内容，如决策点、行动项、责任人、时间节点等，进行总结归纳，然后模仿人工撰写会议纪要的格式及行文风格，以WORD文档形式输出会议纪要。在处理过程中，你需要理解基本的专业术语，将口语化的简称转化为规范的行业术语。
请仔细收听以下会议录音：
<会议录音>
{{MEETING_RECORDING}}
</会议录音>
处理会议录音时，请按照以下步骤进行：
1. 仔细聆听会议录音，**特别关注党委书记和总经理的发言**，记录下所有决策点、行动项、责任人以及时间节点。
2. 识别口语化简称，将其转化为规范的行业术语。
3. **对领导讲话进行重点提炼和总结**，突出指导意义和实际价值。
4. 对记录的信息进行整理和归纳，使内容条理清晰。

撰写会议纪要时，请遵循以下格式及行文风格要求：
1. 标题：简洁明了地概括会议主题。
2. 会议基本信息：包括会议时间、地点、参会人员等。
3. **领导讲话要点**：重点突出党委书记和总经理的重要讲话内容，使用**粗体**标记关键词汇（如**行动项**、**责任人**、**时间节点**等）
4. 会议内容：按照决策点、行动项、责任人、时间节点的顺序依次阐述，使用规范、正式的语言，表达清晰准确。
4. 结尾：可以适当添加总结性话语，如“本次会议达成了[具体成果]，后续需相关责任人按时完成行动项，以推动工作顺利进行”。

**特别注意领导讲话提炼要求：**
1. **重点关注党委书记和总经理的发言内容**，这是会议纪要的核心部分
2. **提炼领导讲话的关键要点**，包括：
   - 重要指示和要求
   - 工作部署和安排
   - 决策事项和原则
   - 强调的重点工作
   - 提出的具体要求和标准
3. **将领导的口语化表达转化为正式的书面语言**
4. **突出领导讲话的权威性和指导性**，使用"要求"、"强调"、"指出"等正式表述

**格式要求：**
- 使用markdown格式，重要内容用**粗体**标记
- 行动项格式：**行动项**：具体内容
- 责任人格式：**责任人**：具体人员
- 时间节点格式：**时间节点**：具体时间

请直接输出完整的会议纪要内容，以符合WORD文档的阅读习惯。内容要丰富、全面，确保涵盖所有关键信息。不要使用任何XML标签包围内容。
"""

    def _build_user_prompt(self, content: str, key_info: Dict[str, Any], meeting_info: Dict[str, Any]) -> str:
        """构建用户提示词"""
        topic = meeting_info.get("topic", "工作会议")
        date = meeting_info.get("date", datetime.now().strftime("%Y年%m月%d日"))
        location = meeting_info.get("location", "公司会议室")
        host = meeting_info.get("host", "党委书记")
        attendees = meeting_info.get("attendees", ["总经理", "相关部门负责人"])

        # 获取民航术语词汇表
        glossary = self._get_aviation_glossary()

        # 获取系统提示词模板并替换占位符
        system_prompt = self._build_system_prompt()

        # 替换模板中的占位符
        user_prompt = system_prompt.replace("{{MEETING_NAME}}", topic)
        user_prompt = user_prompt.replace("{{MEETING_TIME}}", date)
        user_prompt = user_prompt.replace("{{MEETING_LOCATION}}", location)
        user_prompt = user_prompt.replace("{{ATTENDEES}}", ', '.join(attendees))
        user_prompt = user_prompt.replace("{{GLOSSARY}}", glossary)
        user_prompt = user_prompt.replace("{{MEETING_RECORDING}}", content)

        return user_prompt

    def _get_aviation_glossary(self) -> str:
        """获取民航术语词汇表"""
        return """安全管理体系 - SMS (Safety Management System)
飞行品质监控 - FDM (Flight Data Monitoring)
航空器维修 - Aircraft Maintenance
适航管理 - Airworthiness Management
运行控制 - Operations Control
飞行运行 - Flight Operations
机务维修 - Aircraft Engineering
客舱服务 - Cabin Service
地面保障 - Ground Support
航班正常性 - On-time Performance
安全风险评估 - Safety Risk Assessment
持续适航 - Continuing Airworthiness
运行规范 - Operations Specifications
飞行标准 - Flight Standards
维修方案 - Maintenance Program
安全绩效指标 - Safety Performance Indicators
运行监察 - Operations Surveillance
质量管理体系 - QMS (Quality Management System)
应急响应 - Emergency Response
事件调查 - Incident Investigation"""

    def _parse_ai_response(self, response: str) -> Dict[str, Any]:
        """解析AI响应"""
        try:
            # 提取<meeting_minutes>标签内的内容
            meeting_minutes = self._extract_meeting_minutes_content(response.strip())

            # 返回格式化的结果
            return {
                "meeting_summary": meeting_minutes,
                "key_decisions": self._extract_decisions_from_text(meeting_minutes),
                "action_items": self._extract_actions_from_text(meeting_minutes),
                "next_steps": self._extract_next_steps_from_text(meeting_minutes),
                "leadership_remarks": self._extract_leadership_from_text(meeting_minutes)
            }

        except Exception as e:
            logger.error(f"解析AI响应失败: {e}")
            # 返回默认结构
            return {
                "meeting_summary": "AI生成失败，请手动编辑",
                "key_decisions": [],
                "action_items": [],
                "next_steps": [],
                "leadership_remarks": {}
            }

    def _extract_meeting_minutes_content(self, response: str) -> str:
        """提取AI响应中的会议纪要内容"""
        import re

        # 首先尝试提取<meeting_minutes>...</meeting_minutes>标签内的内容（向后兼容）
        match = re.search(r'<meeting_minutes>(.*?)</meeting_minutes>', response, flags=re.DOTALL)

        if match:
            meeting_minutes = match.group(1).strip()
        else:
            # 尝试提取<会议纪要>...</会议纪要>标签内的内容（向后兼容）
            match = re.search(r'<会议纪要>(.*?)</会议纪要>', response, flags=re.DOTALL)
            if match:
                meeting_minutes = match.group(1).strip()
            else:
                # 如果没有找到任何标签，使用整个响应内容
                meeting_minutes = self._filter_thinking_content(response)

        # 清理多余的空行
        meeting_minutes = re.sub(r'\n\s*\n\s*\n', '\n\n', meeting_minutes)

        return meeting_minutes.strip()

    def _filter_thinking_content(self, response: str) -> str:
        """过滤掉AI响应中的<思考>部分内容"""
        import re

        # 移除<思考>...</思考>标签及其内容
        filtered_response = re.sub(r'<思考>.*?</思考>', '', response, flags=re.DOTALL)

        # 移除可能的其他思考标记
        filtered_response = re.sub(r'<思考>.*', '', filtered_response, flags=re.DOTALL)
        filtered_response = re.sub(r'.*</思考>', '', filtered_response, flags=re.DOTALL)

        # 清理多余的空行
        filtered_response = re.sub(r'\n\s*\n\s*\n', '\n\n', filtered_response)

        return filtered_response.strip()

    def _extract_decisions_from_text(self, text: str) -> List[str]:
        """从文本中提取决策事项"""
        decisions = []
        lines = text.split('\n')

        for line in lines:
            line = line.strip()
            if any(keyword in line for keyword in ['决定', '同意', '批准', '通过', '确定']):
                if line and not line.startswith('#') and len(line) > 10:
                    decisions.append(line)

        return decisions[:5]  # 限制数量

    def _extract_actions_from_text(self, text: str) -> List[str]:
        """从文本中提取行动项"""
        actions = []
        lines = text.split('\n')

        for line in lines:
            line = line.strip()
            if any(keyword in line for keyword in ['要求', '需要', '应当', '必须', '负责', '完成']):
                if line and not line.startswith('#') and len(line) > 10:
                    actions.append(line)

        return actions[:8]  # 限制数量

    def _extract_next_steps_from_text(self, text: str) -> List[str]:
        """从文本中提取下一步工作"""
        next_steps = []
        lines = text.split('\n')

        for line in lines:
            line = line.strip()
            if any(keyword in line for keyword in ['下一步', '接下来', '今后', '下阶段', '后续', '下周', '下月']):
                if line and not line.startswith('#') and len(line) > 10:
                    next_steps.append(line)

        return next_steps[:5]  # 限制数量

    def _extract_leadership_from_text(self, text: str) -> Dict[str, str]:
        """从文本中提取领导讲话"""
        leadership = {}
        lines = text.split('\n')
        current_leader = None
        current_content = []

        for line in lines:
            line = line.strip()
            if '党委书记' in line or '总经理' in line:
                # 保存之前的领导讲话
                if current_leader and current_content:
                    leadership[current_leader] = ' '.join(current_content)

                # 开始新的领导讲话
                if '党委书记' in line:
                    current_leader = '党委书记'
                elif '总经理' in line:
                    current_leader = '总经理'
                current_content = [line]
            elif current_leader and line and not line.startswith('#'):
                current_content.append(line)

        # 保存最后一个领导的讲话
        if current_leader and current_content:
            leadership[current_leader] = ' '.join(current_content)

        return leadership

    def _generate_summary(self, content: str, meeting_info: Dict[str, Any]) -> str:
        """生成会议摘要（Markdown格式）"""
        topic = meeting_info.get("topic", "工作会议")
        date = meeting_info.get("date", datetime.now().strftime("%Y年%m月%d日"))
        host = meeting_info.get("host", "党委书记")
        attendees = meeting_info.get("attendees", ["总经理", "相关部门负责人"])

        # 生成高质量的Markdown格式会议纪要
        summary = f"""# {topic}会议纪要

## 会议基本信息

- **会议主题**：{topic}
- **会议时间**：{date}
- **会议地点**：公司会议室
- **主持人**：{host}
- **参会人员**：{', '.join(attendees)}

## 会议内容

### 主要议题

1. **工作进展汇报**
   - 各部门汇报了近期工作完成情况
   - 分析了当前工作中存在的问题和挑战
   - 总结了阶段性工作成果

2. **重点工作讨论**
   - 针对重点项目进行了深入讨论
   - 明确了各部门的职责分工
   - 确定了工作推进的时间节点

3. **下一步工作安排**
   - 制定了下阶段工作计划
   - 明确了工作目标和要求
   - 部署了具体的实施措施

## 会议决议

### 重要决定

- 同意各部门提出的工作方案
- 批准相关项目的推进计划
- 确定了重点工作的实施时间表

### 工作要求

1. 各部门要高度重视，确保工作落实到位
2. 加强部门间的协调配合，形成工作合力
3. 定期汇报工作进展，及时解决遇到的问题

## 领导讲话要点

### {host}讲话

强调要提高政治站位，增强责任意识，确保各项工作任务圆满完成。要求各部门要认真贯彻落实会议精神，以更高的标准、更严的要求推进各项工作。

### 总经理讲话

要求各部门要加强协调配合，统筹推进各项工作。要注重工作质量和效率，确保按时完成既定目标任务。同时要加强风险防控，确保工作安全有序进行。

## 后续行动

1. **本周内**：各部门制定具体实施方案
2. **下周开始**：全面启动相关工作
3. **月底前**：完成阶段性工作目标
4. **持续跟进**：定期召开工作推进会

---

*会议记录：办公室*
*记录时间：{date}*"""

        return summary
    
    def _extract_next_steps(self, content: str) -> List[str]:
        """提取下一步工作"""
        next_steps = []
        
        # 简单的关键词匹配
        keywords = ["下一步", "接下来", "今后", "下阶段", "后续"]
        sentences = content.split("。")
        
        for sentence in sentences:
            if any(keyword in sentence for keyword in keywords):
                next_steps.append(sentence.strip())
        
        return next_steps[:5]  # 限制数量
    
    def _extract_leadership_remarks(self, content: str) -> Dict[str, str]:
        """提取领导讲话要点"""
        # 简化版领导讲话提取
        remarks = {
            "党委书记": "强调要提高政治站位，确保各项工作落到实处。",
            "总经理": "要求各部门加强协调配合，确保完成年度目标任务。"
        }
        
        return remarks
    
    def _format_minutes(self, content: Dict[str, Any], meeting_info: Dict[str, Any]) -> Dict[str, Any]:
        """格式化会议纪要"""
        
        formatted = {
            "title": f"{meeting_info.get('topic', '工作会议')}纪要",
            "header": {
                "meeting_name": meeting_info.get("topic", "工作会议"),
                "date": meeting_info.get("date", datetime.now().strftime("%Y年%m月%d日")),
                "time": meeting_info.get("time", ""),
                "location": meeting_info.get("location", ""),
                "host": meeting_info.get("host", "党委书记"),
                "attendees": meeting_info.get("attendees", ["总经理", "相关部门负责人"]),
            },
            "content": {
                "summary": content["meeting_summary"],
                "decisions": content["key_decisions"],
                "action_items": content["action_items"],
                "next_steps": content["next_steps"],
                "leadership_remarks": content["leadership_remarks"],
            },
            "footer": {
                "recorder": meeting_info.get("recorder", "办公室"),
                "review_date": datetime.now().strftime("%Y年%m月%d日"),
            }
        }
        
        return formatted

    def generate_word_document(self, minutes_data: Dict[str, Any]) -> io.BytesIO:
        """
        生成Word格式的会议纪要文档

        Args:
            minutes_data: 会议纪要数据

        Returns:
            Word文档的字节流
        """
        if not DOCX_AVAILABLE:
            raise Exception("python-docx库未安装，无法生成Word文档")

        # 创建新文档
        doc = Document()

        # 设置默认字体为仿宋
        self._set_document_font(doc, '仿宋')

        # 设置文档标题
        title = doc.add_heading(minutes_data.get('title', '会议纪要'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_paragraph_font(title, '仿宋', 18, bold=True)

        # 添加会议基本信息
        header_info = minutes_data.get('header', {})

        info_heading = doc.add_heading('会议基本信息', level=1)
        self._set_paragraph_font(info_heading, '仿宋', 14, bold=True)

        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'

        info_items = [
            ('会议名称', header_info.get('meeting_name', '')),
            ('会议时间', header_info.get('date', '')),
            ('会议地点', header_info.get('location', '公司会议室')),
            ('主持人', header_info.get('host', '')),
            ('参会人员', ', '.join(header_info.get('attendees', []))),
            ('记录人', header_info.get('recorder', '办公室'))
        ]

        for i, (label, value) in enumerate(info_items):
            label_cell = info_table.cell(i, 0)
            value_cell = info_table.cell(i, 1)

            # 清空单元格内容，然后添加格式化文本
            label_cell.text = ""
            value_cell.text = ""

            # 为标签添加格式化文本（支持粗体）
            self._add_formatted_text_to_cell(label_cell, label)
            # 为值添加格式化文本（支持粗体）
            self._add_formatted_text_to_cell(value_cell, str(value))

        # 添加会议内容
        content = minutes_data.get('content', {})

        # 会议纪要主要内容
        if content.get('summary'):
            content_heading = doc.add_heading('会议内容', level=1)
            self._set_paragraph_font(content_heading, '仿宋', 14, bold=True)

            # 将Markdown转换为格式化的Word内容
            self._add_markdown_content(doc, content['summary'])

        # 决策事项
        if content.get('decisions'):
            decisions_heading = doc.add_heading('会议决议', level=1)
            self._set_paragraph_font(decisions_heading, '仿宋', 14, bold=True)
            for i, decision in enumerate(content['decisions'], 1):
                p = doc.add_paragraph()
                self._add_formatted_text(p, f"{i}. {decision}")

        # 行动项目
        if content.get('action_items'):
            actions_heading = doc.add_heading('后续行动', level=1)
            self._set_paragraph_font(actions_heading, '仿宋', 14, bold=True)
            for i, action in enumerate(content['action_items'], 1):
                p = doc.add_paragraph()
                self._add_formatted_text(p, f"{i}. {action}")

        # 领导讲话
        if content.get('leadership_remarks'):
            leadership_heading = doc.add_heading('领导讲话要点', level=1)
            self._set_paragraph_font(leadership_heading, '仿宋', 14, bold=True)
            for speaker, remarks in content['leadership_remarks'].items():
                speaker_heading = doc.add_heading(f'{speaker}讲话', level=2)
                self._set_paragraph_font(speaker_heading, '仿宋', 13, bold=True)
                p = doc.add_paragraph()
                self._add_formatted_text(p, remarks)

        # 文档结尾
        footer_info = minutes_data.get('footer', {})
        doc.add_paragraph()

        recorder_p = doc.add_paragraph()
        self._add_formatted_text(recorder_p, f"记录人：{footer_info.get('recorder', '办公室')}")

        time_p = doc.add_paragraph()
        self._add_formatted_text(time_p, f"记录时间：{footer_info.get('review_date', '')}")

        # 保存到字节流
        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)

        return doc_stream

    def _set_document_font(self, doc, font_name: str):
        """设置文档默认字体"""
        try:
            from docx.shared import Pt
            from docx.oxml.ns import qn

            # 设置文档默认字体
            style = doc.styles['Normal']
            font = style.font
            font.name = font_name
            font.size = Pt(12)

            # 设置中文字体
            style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        except Exception as e:
            logger.warning(f"设置文档字体失败: {e}")

    def _set_paragraph_font(self, paragraph, font_name: str, size: int = 12, bold: bool = False):
        """设置段落字体"""
        try:
            from docx.shared import Pt
            from docx.oxml.ns import qn

            for run in paragraph.runs:
                font = run.font
                font.name = font_name
                font.size = Pt(size)
                font.bold = bold
                # 设置中文字体
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        except Exception as e:
            logger.warning(f"设置段落字体失败: {e}")

    def _set_cell_font(self, cell, font_name: str, size: int = 12, bold: bool = False):
        """设置表格单元格字体"""
        try:
            from docx.shared import Pt
            from docx.oxml.ns import qn

            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.name = font_name
                    font.size = Pt(size)
                    font.bold = bold
                    # 设置中文字体
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        except Exception as e:
            logger.warning(f"设置单元格字体失败: {e}")

    def _add_markdown_content(self, doc, markdown_text: str):
        """将Markdown内容添加到Word文档，保持格式"""
        import re

        lines = markdown_text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue

            # 检查是否是表格
            if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                # 处理表格
                table_lines = []
                j = i
                while j < len(lines) and '|' in lines[j].strip():
                    table_lines.append(lines[j].strip())
                    j += 1

                if len(table_lines) >= 2:  # 至少需要标题行和分隔行
                    self._add_markdown_table(doc, table_lines)
                    i = j
                    continue

            # 处理标题
            if line.startswith('###'):
                heading = doc.add_heading(level=3)
                self._add_formatted_text(heading, line[3:].strip())
                self._set_paragraph_font(heading, '仿宋', 13, bold=True)
            elif line.startswith('##'):
                heading = doc.add_heading(level=2)
                self._add_formatted_text(heading, line[2:].strip())
                self._set_paragraph_font(heading, '仿宋', 14, bold=True)
            elif line.startswith('#'):
                heading = doc.add_heading(level=1)
                self._add_formatted_text(heading, line[1:].strip())
                self._set_paragraph_font(heading, '仿宋', 15, bold=True)
            # 处理列表项
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(p, line[2:].strip())
            elif re.match(r'^\d+\.\s', line):
                # 提取数字列表的内容，去掉原有的数字编号
                content = re.sub(r'^\d+\.\s*', '', line)
                p = doc.add_paragraph(style='List Number')
                self._add_formatted_text(p, content)
            # 处理普通段落
            else:
                # 处理粗体文本
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)

            i += 1

    def _add_markdown_table(self, doc, table_lines):
        """将Markdown表格添加到Word文档"""
        import re

        # 解析表格数据
        rows = []
        for line in table_lines:
            # 跳过分隔行（包含 --- 的行）
            if re.match(r'^\|[\s\-\|:]+\|$', line.strip()):
                continue

            # 分割单元格，移除首尾的 |
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:  # 确保不是空行
                rows.append(cells)

        if not rows:
            return

        # 创建Word表格
        max_cols = max(len(row) for row in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols)

        # 设置表格样式
        table.style = 'Table Grid'

        # 填充表格数据
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < max_cols:
                    cell = table.cell(row_idx, col_idx)
                    # 清空单元格内容
                    cell.text = ""
                    # 添加格式化文本（支持粗体）
                    self._add_formatted_text_to_cell(cell, cell_data)

                    # 如果是标题行（第一行），设置为粗体
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

        # 添加表格后的空行
        doc.add_paragraph()

    def _add_formatted_text_to_cell(self, cell, text: str):
        """添加格式化文本到表格单元格"""
        import re

        # 清空单元格内容
        for paragraph in cell.paragraphs:
            paragraph.clear()

        # 获取第一个段落
        paragraph = cell.paragraphs[0]

        # 分割文本，处理粗体标记 - 使用非贪婪匹配
        parts = re.split(r'(\*\*[^*]+?\*\*)', text)

        for part in parts:
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                # 粗体文本 - 移除前后的**标记
                bold_text = part[2:-2]
                if bold_text.strip():  # 确保不是空文本
                    run = paragraph.add_run(bold_text)
                    font = run.font
                    font.name = '仿宋'
                    font.bold = True
                    from docx.shared import Pt
                    from docx.oxml.ns import qn
                    font.size = Pt(12)
                    # 设置中文字体
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            elif part.strip():  # 只添加非空文本
                # 普通文本
                run = paragraph.add_run(part)
                font = run.font
                font.name = '仿宋'
                font.bold = False
                from docx.shared import Pt
                from docx.oxml.ns import qn
                font.size = Pt(12)
                # 设置中文字体
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    def _add_formatted_text(self, paragraph, text: str):
        """添加格式化文本到段落"""
        import re

        # 分割文本，处理粗体标记 - 使用非贪婪匹配确保正确处理
        parts = re.split(r'(\*\*[^*]+?\*\*)', text)

        for part in parts:
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                # 粗体文本 - 移除前后的**标记
                bold_text = part[2:-2]
                if bold_text.strip():  # 确保不是空文本
                    run = paragraph.add_run(bold_text)
                    font = run.font
                    font.name = '仿宋'
                    font.bold = True
                    from docx.shared import Pt
                    from docx.oxml.ns import qn
                    font.size = Pt(12)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            elif part.strip():  # 只添加非空文本
                # 普通文本
                run = paragraph.add_run(part)
                font = run.font
                font.name = '仿宋'
                font.bold = False
                from docx.shared import Pt
                from docx.oxml.ns import qn
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')



    def _markdown_to_text(self, markdown_text: str) -> str:
        """将Markdown文本转换为纯文本"""
        import re

        # 移除Markdown标记
        text = re.sub(r'#{1,6}\s*', '', markdown_text)  # 移除标题标记
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)    # 移除粗体标记
        text = re.sub(r'\*(.*?)\*', r'\1', text)        # 移除斜体标记
        text = re.sub(r'`(.*?)`', r'\1', text)          # 移除代码标记
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text) # 移除链接标记
        text = re.sub(r'^[-*+]\s+', '', text, flags=re.MULTILINE)  # 移除列表标记
        text = re.sub(r'^\d+\.\s+', '', text, flags=re.MULTILINE) # 移除有序列表标记

        return text.strip()

    def enhance_content(self, text: str, enhancement_type: str = "expand") -> str:
        """
        内容增强

        Args:
            text: 原始文本
            enhancement_type: 增强类型 (expand/rewrite/optimize)

        Returns:
            增强后的文本
        """
        if self.client:
            return self._ai_enhance_content(text, enhancement_type)
        else:
            # 降级到规则处理
            if enhancement_type == "expand":
                return self._expand_content(text)
            elif enhancement_type == "rewrite":
                return self._rewrite_content(text)
            elif enhancement_type == "optimize":
                return self._optimize_content(text)
            else:
                return text

    def _ai_enhance_content(self, text: str, enhancement_type: str) -> str:
        """使用AI增强内容"""
        try:
            model = os.getenv("ARK_MODEL", "ep-20250511141457-zf4l4")

            # 根据增强类型构建提示词
            prompts = {
                "expand": f"""请对以下会议纪要内容进行扩写，增加必要的细节和说明，使内容更加完整和具体：

原文：
{text}

要求：
1. 保持原意不变
2. 增加具体的实施细节
3. 补充相关的背景信息
4. 使用规范的公文写作风格
5. 字数增加30-50%""",

                "rewrite": f"""请对以下会议纪要内容进行改写，使其更加规范和专业：

原文：
{text}

要求：
1. 使用标准的公文写作语言
2. 规范民航专业术语
3. 优化句式结构
4. 保持内容的准确性
5. 提升表达的专业性""",

                "optimize": f"""请对以下会议纪要内容进行优化，提升其质量和可读性：

原文：
{text}

要求：
1. 消除冗余表达
2. 优化逻辑结构
3. 统一术语使用
4. 增强条理性
5. 保持简洁明了"""
            }

            prompt = prompts.get(enhancement_type, prompts["optimize"])

            response = self.client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "你是一个专业的公文写作助手，擅长民航行业的文档处理。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=1500
            )

            enhanced_text = response.choices[0].message.content.strip()
            logger.info(f"AI内容增强完成: {enhancement_type}")
            return enhanced_text

        except Exception as e:
            error_msg = str(e)
            logger.error(f"AI内容增强失败: {error_msg}")

            # 根据错误类型提供更具体的日志
            if "timeout" in error_msg.lower():
                logger.error("AI内容增强超时，使用本地处理")
            elif "connection" in error_msg.lower():
                logger.error("AI API连接失败，使用本地处理")

            # 降级处理
            logger.info(f"使用本地{enhancement_type}处理")
            if enhancement_type == "expand":
                return self._expand_content(text)
            elif enhancement_type == "rewrite":
                return self._rewrite_content(text)
            else:
                return self._optimize_content(text)
    
    def _expand_content(self, text: str) -> str:
        """扩写内容"""
        # 简化版扩写
        expanded = f"""
        {text}
        
        针对上述内容，需要进一步明确以下几个方面：
        1. 具体实施方案和时间节点
        2. 责任分工和考核标准  
        3. 风险防控和应急预案
        4. 资源保障和支持措施
        """
        return expanded.strip()
    
    def _rewrite_content(self, text: str) -> str:
        """改写内容"""
        # 简化版改写
        return aviation_processor.normalize_text(text)
    
    def _optimize_content(self, text: str) -> str:
        """优化内容"""
        # 简化版优化
        optimized = aviation_processor.normalize_text(text)
        optimized = self._clean_filler_words(optimized)
        return optimized

    def process_text_with_prompt(self, prompt: str) -> str:
        """
        使用自定义提示词处理文本

        Args:
            prompt: 包含处理指令和文本的完整提示词

        Returns:
            处理后的文本
        """
        if self.client:
            return self._ai_process_with_prompt(prompt)
        else:
            # 降级处理：简单返回提示词中的文本部分
            logger.warning("AI服务不可用，返回原始文本")
            # 尝试从提示词中提取原始文本（简单的后备方案）
            lines = prompt.split('\n')
            for i, line in enumerate(lines):
                if '"' in line and i < len(lines) - 1:
                    # 找到引号包围的文本
                    start = line.find('"')
                    if start != -1:
                        end = line.rfind('"')
                        if end > start:
                            return line[start+1:end]
            return "AI服务不可用，无法处理文本"

    def _ai_process_with_prompt(self, prompt: str) -> str:
        """使用AI处理自定义提示词"""
        try:
            model = os.getenv("ARK_MODEL", "ep-20250511141457-zf4l4")

            logger.info(f"使用AI处理自定义提示词，长度: {len(prompt)}")

            # 使用直接API调用
            try:
                response = self._direct_api_call_simple(model, prompt)
                logger.info("AI自定义提示词处理成功")
                return response.strip()
            except Exception as e:
                logger.warning(f"直接API调用失败: {e}，尝试OpenAI客户端")
                # 降级到OpenAI客户端
                response = self.client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "你是一个专业的文本处理助手。请根据用户的要求处理文本，直接返回处理后的结果，不要包含任何解释或说明。"},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3,
                    max_tokens=2000,
                    timeout=60  # 1分钟超时
                )
                result = response.choices[0].message.content.strip()
                logger.info("AI自定义提示词处理成功（OpenAI客户端）")
                return result

        except Exception as e:
            error_msg = str(e)
            logger.error(f"AI自定义提示词处理失败: {error_msg}")

            # 根据错误类型提供更具体的日志
            if "timeout" in error_msg.lower():
                logger.error("AI处理超时")
            elif "connection" in error_msg.lower():
                logger.error("AI API连接失败")
            elif "401" in error_msg or "unauthorized" in error_msg.lower():
                logger.error("AI API认证失败")

            # 返回错误信息
            return f"AI处理失败: {error_msg}"


class MeetingTemplates:
    """会议纪要模板"""
    
    def __init__(self):
        self.templates = {
            "standard": self._get_standard_template(),
            "executive": self._get_executive_template(),
            "safety": self._get_safety_template(),
            "operation": self._get_operation_template(),
        }
    
    def _get_standard_template(self) -> Dict[str, str]:
        """标准会议纪要模板"""
        return {
            "title": "{meeting_name}纪要",
            "header": """
时间：{date} {time}
地点：{location}
主持人：{host}
参会人员：{attendees}
记录人：{recorder}
            """,
            "content": """
一、会议概况
{summary}

二、主要决策事项
{decisions}

三、工作安排
{action_items}

四、下一步工作
{next_steps}

五、领导讲话要点
{leadership_remarks}
            """,
            "footer": """
本纪要已经与会人员确认。

记录人：{recorder}
日期：{review_date}
            """
        }
    
    def _get_executive_template(self) -> Dict[str, str]:
        """高管会议模板"""
        return {
            "title": "四川航空股份有限公司{meeting_name}纪要",
            "header": """
会议时间：{date} {time}
会议地点：{location}  
主持人：{host}
参会人员：{attendees}
            """,
            "content": """
根据公司工作安排，{date}召开{meeting_name}。现将会议主要内容纪要如下：

一、会议议题
{summary}

二、决策事项
{decisions}

三、工作部署
{action_items}

四、领导指示
{leadership_remarks}

五、下步安排
{next_steps}
            """
        }
    
    def _get_safety_template(self) -> Dict[str, str]:
        """安全会议模板"""
        return self._get_standard_template()
    
    def _get_operation_template(self) -> Dict[str, str]:
        """运营会议模板"""
        return self._get_standard_template()
    
    def get_template(self, template_type: str = "standard") -> Dict[str, str]:
        """获取模板"""
        return self.templates.get(template_type, self.templates["standard"])

    def _chunked_ai_generate(self, content: str, key_info: Dict[str, Any], meeting_info: Dict[str, Any], model: str) -> Dict[str, Any]:
        """分段处理长内容的AI生成"""
        try:
            logger.info("开始分段处理长内容")

            # 分割内容
            chunks = self._split_content_by_sentences(content, max_length=6000)
            logger.info(f"内容分割为{len(chunks)}段")

            # 生成各段摘要
            summaries = []
            for i, chunk in enumerate(chunks):
                logger.info(f"正在处理第{i+1}/{len(chunks)}段")
                try:
                    summary = self._generate_chunk_summary(chunk, i, len(chunks), meeting_info, model)
                    summaries.append(summary)
                except Exception as e:
                    logger.error(f"第{i+1}段处理失败: {e}")
                    summaries.append(f"第{i+1}段处理失败: {str(e)}")

            logger.info("所有分段摘要生成完成，开始生成最终纪要")

            # 基于摘要生成最终纪要
            combined_summary = "\n\n".join([f"## 分段{i+1}摘要\n{summary}" for i, summary in enumerate(summaries)])
            final_result = self._generate_final_minutes_from_summaries(combined_summary, key_info, meeting_info, model)

            logger.info("分段处理完成")
            return final_result

        except Exception as e:
            logger.error(f"分段处理失败: {e}")
            # 降级到本地生成
            return self._generate_local_content(content, key_info, meeting_info)

    def _split_content_by_sentences(self, content: str, max_length: int = 6000) -> List[str]:
        """按句子分割内容"""
        # 按句号分割
        sentences = content.split('。')
        chunks = []
        current_chunk = ""

        for sentence in sentences:
            # 检查添加这个句子后是否超过长度限制
            test_chunk = current_chunk + sentence + '。'
            if len(test_chunk) > max_length and current_chunk:
                # 当前块已满，保存并开始新块
                chunks.append(current_chunk.strip())
                current_chunk = sentence + '。'
            else:
                current_chunk = test_chunk

        # 添加最后一块
        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        return chunks

    def _generate_chunk_summary(self, chunk: str, chunk_index: int, total_chunks: int, meeting_info: Dict[str, Any], model: str) -> str:
        """生成单个分段的摘要"""
        try:
            topic = meeting_info.get("topic", "工作会议")

            prompt = f"""
请对以下会议记录片段进行总结，提取关键信息：

会议主题：{topic}
片段：{chunk_index + 1}/{total_chunks}

会议记录片段：
{chunk}

请提取以下信息：
1. 主要讨论议题
2. 重要决策和结论
3. 关键数据和指标
4. 行动计划和责任人
5. 其他重要信息

请用简洁的要点形式总结，保持客观准确。
"""

            # 使用直接API调用
            try:
                response = self._direct_api_call_simple(model, prompt)
                return response.strip()
            except Exception as e:
                logger.warning(f"直接API调用失败: {e}，尝试OpenAI客户端")
                # 降级到OpenAI客户端
                response = self.client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=1000,
                    temperature=0.3,
                    timeout=300  # 5分钟超时
                )
                return response.choices[0].message.content.strip()

        except Exception as e:
            logger.error(f"生成分段摘要失败: {e}")
            return f"分段{chunk_index + 1}摘要生成失败: {str(e)}"

    def _direct_api_call_simple(self, model: str, prompt: str) -> str:
        """简化的直接API调用"""
        ark_api_key = os.getenv("ARK_API_KEY")
        ark_base_url = os.getenv("ARK_BASE_URL", "https://ark.cn-beijing.volces.com/api/v3")

        url = f"{ark_base_url}/chat/completions"
        headers = {
            "Authorization": f"Bearer {ark_api_key}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": model,
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.3,
            "max_tokens": 1000
        }

        timeout = httpx.Timeout(timeout=300.0)  # 5分钟超时

        with httpx.Client(timeout=timeout) as client:
            response = client.post(url, headers=headers, json=payload)
            response.raise_for_status()
            result = response.json()
            return result["choices"][0]["message"]["content"]

    def _generate_final_minutes_from_summaries(self, combined_summary: str, key_info: Dict[str, Any], meeting_info: Dict[str, Any], model: str) -> Dict[str, Any]:
        """基于分段摘要生成最终会议纪要"""
        try:
            # 构建系统提示词
            system_prompt = self._build_system_prompt()

            # 构建用户提示词（使用合并的摘要）
            user_prompt = self._build_user_prompt(combined_summary, key_info, meeting_info)

            # 调用AI生成最终纪要
            try:
                ai_response = self._direct_api_call(model, system_prompt, user_prompt)
            except Exception as e:
                logger.warning(f"直接API调用失败: {e}，尝试OpenAI客户端")
                response = self.client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=0.3,
                    max_tokens=3000,
                    timeout=600  # 10分钟超时
                )
                ai_response = response.choices[0].message.content

            return self._parse_ai_response(ai_response)

        except Exception as e:
            logger.error(f"生成最终会议纪要失败: {e}")
            # 降级到本地生成
            return self._generate_local_content(combined_summary, key_info, meeting_info)
